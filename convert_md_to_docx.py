#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将Markdown文件转换为Word文档，适用于毕业论文格式
"""

import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
import re


def convert_md_to_docx(md_path, docx_path):
    """
    将Markdown文件转换为Word文档
    :param md_path: Markdown文件路径
    :param docx_path: 输出Word文档路径
    """
    # 读取Markdown文件
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 转换Markdown为HTML
    html_content = markdown.markdown(md_content, extensions=['extra', 'toc', 'codehilite'])
    
    # 解析HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # 创建Word文档
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.5)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
    
    # 处理标题和内容
    handle_element(soup, doc)
    
    # 保存Word文档
    doc.save(docx_path)
    print(f"转换完成！文件已保存至: {docx_path}")


def handle_element(element, doc):
    """
    递归处理HTML元素
    :param element: HTML元素
    :param doc: Word文档对象
    """
    if element.name == 'h1':
        # 一级标题
        para = doc.add_heading(element.get_text(), level=1)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(12)
    elif element.name == 'h2':
        # 二级标题
        para = doc.add_heading(element.get_text(), level=2)
        para.paragraph_format.space_after = Pt(6)
    elif element.name == 'h3':
        # 三级标题
        para = doc.add_heading(element.get_text(), level=3)
        para.paragraph_format.space_after = Pt(4)
    elif element.name == 'h4':
        # 四级标题
        para = doc.add_heading(element.get_text(), level=4)
        para.paragraph_format.space_after = Pt(2)
    elif element.name == 'p':
        # 段落
        text = element.get_text()
        if text.strip():
            para = doc.add_paragraph(text)
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.space_after = Pt(6)
    elif element.name == 'ul' or element.name == 'ol':
        # 列表
        for li in element.find_all('li', recursive=False):
            text = li.get_text()
            if element.name == 'ul':
                doc.add_paragraph(text, style='List Bullet')
            else:
                doc.add_paragraph(text, style='List Number')
    elif element.name == 'table':
        # 表格
        table = doc.add_table(rows=0, cols=0)
        # 处理表头
        thead = element.find('thead')
        if thead:
            header_row = table.add_row()
            for th in thead.find_all('th'):
                cell = header_row.add_cell()
                cell.text = th.get_text()
                # 设置表头样式
                for paragraph in cell.paragraphs:
                    paragraph.style = 'Table Header'
        # 处理表体
        tbody = element.find('tbody')
        if tbody:
            for tr in tbody.find_all('tr'):
                row = table.add_row()
                for td in tr.find_all('td'):
                    cell = row.add_cell()
                    cell.text = td.get_text()
        # 自动调整表格列宽
        table.autofit = True
    elif element.name == 'pre':
        # 代码块
        code = element.find('code')
        if code:
            para = doc.add_paragraph()
            run = para.add_run(code.get_text())
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            para.paragraph_format.space_after = Pt(6)
    elif element.name == 'hr':
        # 分隔线
        doc.add_paragraph().add_run('\n' + '—' * 40 + '\n')
    elif element.name == 'div':
        # 处理div元素
        for child in element.children:
            handle_element(child, doc)
    elif element.name == 'body':
        # 处理body元素的子元素
        for child in element.children:
            handle_element(child, doc)


if __name__ == '__main__':
    # 输入和输出文件路径
    md_file = '论文.md'
    docx_file = '毕业论文_转换版.docx'
    
    # 执行转换
    convert_md_to_docx(md_file, docx_file)
